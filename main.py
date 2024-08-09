import os
import re
from functions import extract_text_from_pdf,split_text_by_heading,summarize_chapters_replicate
from docx import Document
os.environ["REPLICATE_API_TOKEN"] = "r8_OYXuIDs1GE9zPtbVFXmdw1KCkXaNTq64QOeAG"




def save_text_to_file(text, output_path):
    # Ensure the directory exists
    output_dir = os.path.dirname(output_path)
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    with open(output_path, 'w', encoding='utf-8') as file:
        file.write(text)
    print(f"Entire text saved to {output_path}")

def save_chapters_to_files(chapters, output_dir):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    for chapter_name, chapter_text in chapters.items():
        # Sanitize filename
        safe_title = re.sub(r'[\/:*?"<>|]', '', chapter_name)
        file_path = os.path.join(output_dir, f"{safe_title}.txt")
        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(chapter_text)
    
    print(f"Chapters saved to {output_dir}")
    


# Function to summarize all chapters in a folder
def summarize_chapters_in_folder(folder_path, summaries_folder_path):
    # Ensure the summaries folder exists
    if not os.path.exists(summaries_folder_path):
        os.makedirs(summaries_folder_path)
    
    # Iterate through each file in the folder
    for filename in os.listdir(folder_path):
        if filename.endswith(".txt"):
            chapter_path = os.path.join(folder_path, filename)
            
            # Read the chapter text
            with open(chapter_path, 'r', encoding='utf-8') as file:
                chapter_text = file.read()
            
            # Get the summary
            summary = summarize_chapters_replicate.summarize_chapters(chapter_text)
            summary = ''.join(summary)
            # Save the summary to a new file
            summary_filename = f"summary_{filename}"
            summary_path = os.path.join(summaries_folder_path, summary_filename)
            with open(summary_path, 'w', encoding='utf-8') as file:
                file.write(summary)
    
                
def create_word_document_from_summaries(summaries_folder_path, output_docx_path):
    # Create a new Document
    doc = Document()
    doc.add_heading('Summaries', level=1)
    titles = list(chapters.values())
    title = []
    for each in titles:
        title.append(each.splitlines()[0])
    count=0
    # Iterate through each file in the summaries folder
    for filename in os.listdir(summaries_folder_path):
        
        if filename.startswith('summary_') and filename.endswith('.txt'):
            summary_path = os.path.join(summaries_folder_path, filename)
            
            # Read the summary text
            with open(summary_path, 'r', encoding='utf-8') as file:
                summary_text = file.read()
            
            # Add the filename as a heading and the summary as a paragraph
            doc.add_heading(title[count], level=2)
            doc.add_paragraph(summary_text)
            count+=1
    # Save the document
    doc.save(output_docx_path)

# Example usage
pdf_path = '3chapters.pdf'
output_dir = 'output/chapters'
summaries_dir = 'output/summaries'
full_text_path = 'output/full_text.txt'
output_docx = 'output//Summaries.docx'

text = extract_text_from_pdf.extract_text_from_pdf(pdf_path)
save_text_to_file(text, full_text_path)

chapters = split_text_by_heading.split_text_by_headings(text)
save_chapters_to_files(chapters, output_dir)
summarize_chapters_in_folder(output_dir, summaries_dir)
create_word_document_from_summaries(summaries_dir, output_docx)
